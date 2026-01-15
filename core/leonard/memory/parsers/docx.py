"""Parser for Microsoft Word documents."""

from pathlib import Path

from leonard.memory.parsers.base import DocumentParser, ParsedDocument, DocumentSection
from leonard.utils.logging import logger


class DocxParser(DocumentParser):
    """
    Parser for Microsoft Word documents using python-docx.

    Extracts text content preserving paragraph structure.
    Supports: .docx
    """

    @property
    def supported_extensions(self) -> list[str]:
        return [".docx"]

    async def parse(self, file_path: Path) -> ParsedDocument:
        """
        Parse a Word document.

        Args:
            file_path: Path to the .docx file

        Returns:
            ParsedDocument with extracted text content
        """
        self._validate_file(file_path)

        metadata = self._get_base_metadata(file_path)
        metadata["language"] = "docx"

        try:
            from docx import Document

            doc = Document(str(file_path))

            # Extract paragraphs
            paragraphs = []
            sections = []
            current_heading = None
            current_content = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # Check if this is a heading
                if para.style and para.style.name.startswith("Heading"):
                    # Save previous section
                    if current_heading or current_content:
                        sections.append(
                            DocumentSection(
                                title=current_heading,
                                content="\n".join(current_content),
                                level=self._get_heading_level(para.style.name),
                            )
                        )

                    current_heading = text
                    current_content = []
                else:
                    paragraphs.append(text)
                    current_content.append(text)

            # Don't forget the last section
            if current_heading or current_content:
                sections.append(
                    DocumentSection(
                        title=current_heading,
                        content="\n".join(current_content),
                        level=0,
                    )
                )

            content = "\n\n".join(paragraphs)

            # Extract tables (as text)
            for table in doc.tables:
                table_text = self._extract_table_text(table)
                if table_text:
                    content += f"\n\n{table_text}"

            # Document metadata
            if doc.core_properties:
                if doc.core_properties.title:
                    metadata["title"] = doc.core_properties.title
                if doc.core_properties.author:
                    metadata["author"] = doc.core_properties.author
                if doc.core_properties.subject:
                    metadata["subject"] = doc.core_properties.subject

            metadata["paragraph_count"] = len(paragraphs)

            logger.info(f"Parsed DOCX: {file_path.name} ({len(paragraphs)} paragraphs)")

            return ParsedDocument(
                content=content,
                metadata=metadata,
                sections=sections,
            )

        except ImportError:
            raise RuntimeError(
                "python-docx is required for Word document parsing. "
                "Install with: pip install python-docx"
            )
        except Exception as e:
            logger.error(f"Failed to parse DOCX {file_path}: {e}")
            raise RuntimeError(f"Failed to parse Word document: {e}")

    def _get_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name."""
        # Heading 1, Heading 2, etc.
        try:
            return int(style_name.split()[-1])
        except (ValueError, IndexError):
            return 0

    def _extract_table_text(self, table) -> str:
        """Extract text from a table as formatted text."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))

        if rows:
            return "Table:\n" + "\n".join(rows)
        return ""
